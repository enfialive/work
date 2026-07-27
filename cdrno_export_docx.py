#!/usr/bin/env python3
"""CDRNO — FASTA label mode: export full SEQ ID NO list + annotation tables to Word (.docx)"""
import sys, re
sys.path.insert(0, '.')
from cdr_batch import *
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Raw input as FASTA (label + sequence) ──
raw = r""">M1 VH
QVQLQQSGAELVRPGSSVKISCKASGYAFSSYWMNWVKQRPGQGLEWIGQIYPGDGDARYNGKFNGKATLTTDKSSSTAYMQLSSLTSEDSAVYFCTRSMGLGLDYWGQGTTLTVSS
>M1 VL
DVVMTQTPLSLPVSLGHQASISCRSSQSLVHSNGNTYLHWYLQKPGQSPKLLIYKVSNRFSGVPDRFSGSGSGTDFTLKISRVEAEDLGVYFCSQTTHVPPYTFGGGTKLEIK
>M2 VH
QVTLKESGPGILQSSQTLSLTCSFSGFSLSTSGMGVSWIRQPSGKGLEWLAHIYWDDDKRYNPSLKSRLTISKDTSRNQVFLKITSVDTADTATYYCARRAPYFDSPFAYWGQGTLVTVSA
>M2 VL
NIMMTQSPSSLAVSAGEKVTMSCKSSQSVLYSSNQKNYLAWYQQKPGQSPKLLIYWASTRESGVPDRFTGSGSGTDFTLTISSVQAEDLAVYYCHQYLSSLTFGAGTMLELK
>M3 VH
QVTLKESGPGILQSSQTLSLTCSFSGFSLSTSGMGVSWIRQPSGKGLEWLAHIYWDDDKRYNPSLKSRLTISKDTSRNQVFLKITSVDTADTATYYCARTIPTVVFDYWGQGTTLTVSS
>M4 VH
DIVMSQSPSSLAVSVGEKVTMSCKSSQSLLYSSNQKNYLAWYQQKPGQSPKLLIYWASTRESGVPDRFTGSGSGTDFTLTISSVKAEDLAVYYCQQYYSYRTFGGGTKLEIK
>M4 VH
QVTLKESGPGILQSSQTLSLTCSFSGFSLSTSGMGVSWIRQPSGKGLEWLAHIYWDDDKRYNPSLKSRLTISKDTSRNQVFLKITSVDTADTATYYCARTIPTVVFDYWGQGTTLTVSS
>M4 VL
DILMTQSPSSMSVSLGDAVSITCHASQVISSNIGWLQQKPGKSFKGLIYHGTNLEDGVPSRFSGSGSGADYSLTISSLESEDFADYYCVRYAQFPWTFGGGTKLEIK
>M5 VH
QVTLKESGPGILQSSQTLSLTCSFSGFSLSTSGMGVSWIRQPSGKGLEWLAHIYWDDDKRYNPSLKSRLTISKDTSRNQVFLRITSVDTADSATYHCARTTTTVVFDYWGQGTTLTVSS
>M5 VL
DIVMSQSPSSLAVSVGEKVTMSCKSSQSLLYSSNQKNYLAWYQQKPGQSPKLLIYWASTRESGVPDRFTGSGSGTDFTLTISSVKAEDLAVYYCQQYYSYRTFGGGTKLEIK"""

# ── Parse FASTA ──
parsed = []
dup_counts = {}
for block in raw.strip().split('>'):
    if not block.strip():
        continue
    lines = block.strip().split('\n')
    raw_label = lines[0].strip()
    seq = ''.join(l.strip() for l in lines[1:]).upper()
    seq = re.sub(r'[^ACDEFGHIKLMNPQRSTVWY]', '', seq)

    # Handle duplicate labels
    if raw_label in dup_counts:
        dup_counts[raw_label] += 1
        label = f"{raw_label} ({dup_counts[raw_label]})"
    else:
        dup_counts[raw_label] = 1
        label = raw_label

    parsed.append((raw_label, label, seq))

# ── Auto-detect chain type ──
detected_ct = []
for raw_label, label, s in parsed:
    ct = detect_chain(s)
    detected_ct.append(ct)

schemes = ['Kabat', 'Chothia', 'IMGT', 'AbM']
chain_full = {'VH': 'Heavy', 'VK': 'VL', 'VL': 'VL'}
cdr_names_map = {
    'VH': ['CDR-H1', 'CDR-H2', 'CDR-H3'],
    'VK': ['CDR-L1', 'CDR-L2', 'CDR-L3'],
    'VL': ['CDR-L1', 'CDR-L2', 'CDR-L3'],
}

# ── SEQ ID NO assignments ──
vseq_to_id = {}
vseq_entries = []  # (sid, description, seq)
next_id = 1

for (raw_label, label, s), ct in zip(parsed, detected_ct):
    if s not in vseq_to_id:
        vseq_to_id[s] = next_id
        next_id += 1

vseq_id_list = []  # (sid, raw_label, label, ct, seq)
for (raw_label, label, s), ct in zip(parsed, detected_ct):
    vseq_id_list.append((vseq_to_id[s], raw_label, label, ct, s))

# CDR collection
cdr_raw = []  # (label, ct, cdr_name, scheme, seq)
for (raw_label, label, s), ct in zip(parsed, detected_ct):
    for d in schemes:
        result = analyze_cdr(s, d)
        for reg in result['regions']:
            if reg['name'].startswith('CDR'):
                cdr_raw.append((label, ct, reg['name'], d, reg['sequence']))

cdr_seq_to_id = {}
for label, ct, cdr_name, scheme, seq in cdr_raw:
    if seq not in cdr_seq_to_id:
        cdr_seq_to_id[seq] = next_id
        next_id += 1

cdr_id_list = []  # (sid, label, ct, cdr_name, scheme, seq)
for label, ct, cdr_name, scheme, seq in cdr_raw:
    cdr_id_list.append((cdr_seq_to_id[seq], label, ct, cdr_name, scheme, seq))

# ── Build all_unique list for Table 2 ──
all_unique = []

# V-region: group by seq to merge descriptions
vseq_meta = {}  # seq -> (sid, list of descriptions)
for sid, raw_label, label, ct, seq in vseq_id_list:
    desc = f"{raw_label} [{chain_full.get(ct, ct)}]"
    if seq not in vseq_meta:
        vseq_meta[seq] = (sid, [])
    vseq_meta[seq][1].append(desc)

for seq, (sid, descs) in vseq_meta.items():
    all_unique.append((sid, 'V-region', seq, ', '.join(descs)))

for seq, sid in sorted(cdr_seq_to_id.items(), key=lambda x: x[1]):
    # Collect all (variant, scheme) sources for this CDR sequence
    sources = []
    seen_sources = set()
    cdr_type = 'CDR'
    for s, lb, ct2, cn, sc, seq2 in cdr_id_list:
        if seq2 == seq:
            cdr_type = cn  # CDR-H1 / CDR-L2 etc.
            key = (lb, sc)
            if key not in seen_sources:
                sources.append(key)
                seen_sources.add(key)
    # Build description: e.g. "M1 VH (Kabat), M2 VH (Kabat)"
    desc_parts = [f'{v} ({s})' for v, s in sources]
    desc = ', '.join(desc_parts)
    all_unique.append((sid, cdr_type, seq, desc))

all_unique.sort(key=lambda x: x[0])

# ── Build table 1A (VH) and 1B (VL) data ──
def build_cdr_table(chain_type):
    """Return rows: [(variant_label, [(cdr_name, {scheme:sid}, ct)])]"""
    variants = []
    for (raw_label, label, s), ct in zip(parsed, detected_ct):
        if ct == chain_type:
            variants.append(label)
    if not variants:
        return []

    cdr_names = cdr_names_map[chain_type]
    rows = []
    ref_variant = variants[0]
    for var in variants:
        cell_data = {}
        for cn in cdr_names:
            cell_data[cn] = {}
            for d in schemes:
                matches = [s for s, lb, ct2, cn2, sc, seq in cdr_id_list
                           if lb == var and cn2 == cn and sc == d]
                sid = matches[0] if matches else None
                ref_matches = [s for s, lb, ct2, cn2, sc, seq in cdr_id_list
                               if lb == ref_variant and cn2 == cn and sc == d]
                ref_sid = ref_matches[0] if ref_matches else None
                cell_data[cn][d] = (sid, sid != ref_sid if var != ref_variant else False)
        ct_of_var = detected_ct[[i for i, (rl, l, s) in enumerate(parsed)].index(
            [j for j, (rl, l, s) in enumerate(parsed) if l == var][0])]
        rows.append((var, cell_data, ct_of_var))
    return rows

# ── Generate DOCX ──
doc = Document()
section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(1.2)
section.right_margin = Cm(1.2)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ── Title ──
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_para.add_run('CDRNO Analysis — SEQ ID NO Sequence List')
r.bold = True
r.font.size = Pt(14)

info = doc.add_paragraph()
r = info.add_run(
    f'Total unique: {len(all_unique)} (V-region: {len(vseq_to_id)}, CDR: {len(cdr_seq_to_id)})  |  '
    f'Input: {len(parsed)} sequences ({sum(1 for c in detected_ct if c=="VH")} VH + '
    f'{sum(1 for c in detected_ct if c=="VK")} VK + {sum(1 for c in detected_ct if c=="VL")} VL)  |  '
    f'Schemes: Kabat / Chothia / IMGT / AbM'
)
r.font.size = Pt(9)

doc.add_paragraph()

# ── Table 0: V-region ──
t0_heading = doc.add_paragraph()
r = t0_heading.add_run('Table 0: V-region Sequences')
r.bold = True
r.font.size = Pt(11)

t0 = doc.add_table(rows=1, cols=5)
t0.style = 'Table Grid'
t0.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (text, w) in enumerate(zip(['#', 'User Label', 'Detected Chain', 'Length', 'SEQ ID NO'],
                                    [Cm(1.0), Cm(3.5), Cm(3.0), Cm(1.5), Cm(2.5)])):
    t0.rows[0].cells[i].width = w
    p = t0.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(255, 255, 255)
    shd = t0.rows[0].cells[i]._element.get_or_add_tcPr()
    shd.append(shd.makeelement(qn('w:shd'), {qn('w:fill'): '2F5496', qn('w:val'): 'clear'}))

for idx, (sid, raw_label, label, ct, seq) in enumerate(vseq_id_list):
    row = t0.add_row()
    vals = [str(idx + 1), f"{raw_label} ({label})" if label != raw_label else raw_label,
            f"{chain_full.get(ct, ct)} ({ct})", str(len(seq)), f"SEQ ID NO: {sid}"]
    for ci, v in enumerate(vals):
        p = row.cells[ci].paragraphs[0]
        if ci in (0, 3):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(v)
        r.font.size = Pt(8)
        if ci == 1 and ct != ('VH' if 'VH' in raw_label else ('VK' if 'VL' not in raw_label else 'VL')):
            # Detect mismatch: label says VH but chain detected as VL/VK
            expected_ct = 'VH' if 'VH' in raw_label else ('VL' if 'VL' in raw_label else 'VK')
            if (expected_ct == 'VH' and ct != 'VH') or (expected_ct in ('VL','VK') and ct not in ('VK','VL')):
                r.font.color.rgb = RGBColor(200, 0, 0)

doc.add_paragraph()

# ── Chain type mismatch warnings ──
mismatches = []
for (raw_label, label, s), ct in zip(parsed, detected_ct):
    expected = 'VH' if 'VH' in raw_label else ('VL' if 'VL' in raw_label else 'VK')
    if (expected == 'VH' and ct != 'VH') or (expected in ('VL', 'VK') and ct not in ('VK', 'VL')):
        mismatches.append((raw_label, expected, ct))

if mismatches:
    warn = doc.add_paragraph()
    r = warn.add_run('[!] Chain type mismatches detected (label vs auto-detection):')
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(200, 0, 0)
    for ml, exp, got in mismatches:
        line = doc.add_paragraph()
        r = line.add_run(f'    {ml}: label suggests {exp}, but sequence detected as {chain_full.get(got, got)} ({got}). Using auto-detected chain for CDR naming.')
        r.font.size = Pt(8); r.font.color.rgb = RGBColor(180, 0, 0)

doc.add_paragraph()

# ── Table 1A: VH CDR ──
vh_rows = build_cdr_table('VH')
vl_rows = build_cdr_table('VK') + build_cdr_table('VL')
vl_cts = sorted(set(r[2] for r in vl_rows)) if vl_rows else []
vl_label = 'Light Chain (' + '/'.join(vl_cts) + ')' if vl_cts else 'Light Chain'

for tbl_rows, chain_label, cdr_names in [
    (vh_rows, 'Heavy Chain (VH)', ['CDR-H1', 'CDR-H2', 'CDR-H3']),
    (vl_rows, vl_label, ['CDR-L1', 'CDR-L2', 'CDR-L3']),
]:
    if not tbl_rows:
        continue

    heading = doc.add_paragraph()
    r = heading.add_run(f'Table 1: CDR Annotation — {chain_label}')
    r.bold = True; r.font.size = Pt(11)

    t1 = doc.add_table(rows=1, cols=7)
    t1.style = 'Table Grid'
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Variant', 'CDR'] + schemes
    widths = [Cm(2.5), Cm(2.0)] + [Cm(3.5)] * 5
    for i, (text, w) in enumerate(zip(headers, widths)):
        t1.rows[0].cells[i].width = w
        p = t1.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(255, 255, 255)
        shd = t1.rows[0].cells[i]._element.get_or_add_tcPr()
        shd.append(shd.makeelement(qn('w:shd'), {qn('w:fill'): '2F5496', qn('w:val'): 'clear'}))

    for (var, cell_data, ct) in tbl_rows:
        for cn in cdr_names:
            row = t1.add_row()
            p0 = row.cells[0].paragraphs[0]
            r0 = p0.add_run(var)
            r0.font.size = Pt(8)
            p1 = row.cells[1].paragraphs[0]
            r1 = p1.add_run(cn)
            r1.font.size = Pt(8)

            for di, d in enumerate(schemes):
                sid, is_diff = cell_data[cn][d]
                p = row.cells[di + 2].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f'SEQ ID NO: {sid}' if sid else '—')
                r.font.size = Pt(8)
                if is_diff:
                    r.bold = True
                    r.font.color.rgb = RGBColor(180, 0, 0)

    doc.add_paragraph()

# ── Table 2: Full SEQ ID NO List ──
heading2 = doc.add_paragraph()
r = heading2.add_run(f'Table 2: Complete SEQ ID NO Sequence List ({len(all_unique)} entries)')
r.bold = True; r.font.size = Pt(11)

t2 = doc.add_table(rows=1, cols=5)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_texts = ['SEQ ID NO', 'Type', 'Sequence', 'Length', 'Source Variant(s)']
hdr_widths = [Cm(2.2), Cm(2.2), Cm(14.0), Cm(1.3), Cm(8.5)]
for i, (text, w) in enumerate(zip(hdr_texts, hdr_widths)):
    t2.rows[0].cells[i].width = w
    p = t2.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(255, 255, 255)
    shd = t2.rows[0].cells[i]._element.get_or_add_tcPr()
    shd.append(shd.makeelement(qn('w:shd'), {qn('w:fill'): '2F5496', qn('w:val'): 'clear'}))

for sid, typ, seq, desc in all_unique:
    row = t2.add_row()
    for i, w in enumerate(hdr_widths):
        row.cells[i].width = w

    vals = [
        str(sid),
        typ,
        seq,
        str(len(seq)),
        desc,
    ]
    for ci, v in enumerate(vals):
        p = row.cells[ci].paragraphs[0]
        if ci in (0, 3):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(v)
        if ci == 2:
            r.font.name = 'Consolas'
            r.font.size = Pt(7)
        else:
            r.font.size = Pt(8)

    # Blue background for V-region
    if typ == 'V-region':
        for ci in range(5):
            s_elm = row.cells[ci]._element.get_or_add_tcPr()
            s_shd = s_elm.makeelement(qn('w:shd'), {qn('w:fill'): 'D6E4F0', qn('w:val'): 'clear'})
            s_elm.append(s_shd)

# ── Save ──
output_path = r'C:\Users\admin\Desktop\CC\work\CDRNO_SEQ_ID_NO_List.docx'
doc.save(output_path)
print(f'Done: {output_path}')
print(f'V-region: {len(vseq_to_id)} unique  |  CDR: {len(cdr_seq_to_id)} unique  |  Total: {len(all_unique)}')

# Also print summary to stdout
print('\n=== Table 0: V-region ===')
for idx, (sid, raw_label, label, ct, seq) in enumerate(vseq_id_list):
    mismatch_flag = ''
    expected = 'VH' if 'VH' in raw_label else ('VL' if 'VL' in raw_label else 'VK')
    if (expected == 'VH' and ct != 'VH') or (expected in ('VL', 'VK') and ct not in ('VK', 'VL')):
        mismatch_flag = ' [!] label says ' + expected + ', detected ' + ct
    print(f'  {idx+1}. {label} -> {chain_full.get(ct,ct)} ({ct}) | {len(seq)} aa | SEQ ID NO: {sid}{mismatch_flag}')

print(f'\n=== Table 1A: VH CDR ({len(vh_rows)} variants) ===')
for var, cell_data, ct in vh_rows:
    print(f'  {var}:')
    for cn in ['CDR-H1', 'CDR-H2', 'CDR-H3']:
        parts = []
        for d in schemes:
            sid, diff = cell_data[cn][d]
            parts.append(f'{d}=SEQ ID NO:{sid}{"*" if diff else ""}')
        sep = ' | '
        print(f'    {cn}: {sep.join(parts)}')

print(f'\n=== Table 1B: VL CDR ({len(vl_rows)} variants) ===')
for var, cell_data, ct in vl_rows:
    print(f'  {var}:')
    for cn in ['CDR-L1', 'CDR-L2', 'CDR-L3']:
        parts = []
        for d in schemes:
            sid, diff = cell_data[cn][d]
            parts.append(f'{d}=SEQ ID NO:{sid}{"*" if diff else ""}')
        sep = ' | '
        print(f'    {cn}: {sep.join(parts)}')
