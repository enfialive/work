#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CDRNO — Generic Word export tool
Reads FASTA from stdin or file, generates standardized CDRNO analysis Word document.

Usage:
  python cdrno_generate_docx.py < input.fasta
  python cdrno_generate_docx.py input.fasta
  echo ">M1 VH\nEVQLLES..." | python cdrno_generate_docx.py

Output: CDRNO_SEQ_ID_NO_List.docx in current directory.
"""

import sys, re, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from cdr_batch import analyze_cdr, detect_chain

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Constants ──
SCHEMES = ['Kabat', 'Chothia', 'IMGT', 'AbM']

# ── Parse FASTA ──
def parse_fasta(text):
    """Parse FASTA text, return list of (raw_label, seq_clean)."""
    parsed = []
    dup_counts = {}
    for block in text.strip().split('>'):
        if not block.strip():
            continue
        lines = block.strip().split('\n')
        raw_label = lines[0].strip()
        seq = ''.join(l.strip() for l in lines[1:]).upper()
        seq = re.sub(r'[^ACDEFGHIKLMNPQRSTVWY]', '', seq)

        if raw_label in dup_counts:
            dup_counts[raw_label] += 1
            label = f"{raw_label} ({dup_counts[raw_label]})"
        else:
            dup_counts[raw_label] = 1
            label = raw_label

        parsed.append((raw_label, label, seq))
    return parsed

# ── Helper: styled header row ──
def add_header_row(table, texts, widths):
    for i, (txt, w) in enumerate(zip(texts, widths)):
        table.rows[0].cells[i].width = w
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(255, 255, 255)
        shd = table.rows[0].cells[i]._element.get_or_add_tcPr()
        shd.append(shd.makeelement(qn('w:shd'), {qn('w:fill'): '2F5496', qn('w:val'): 'clear'}))

# ── Main ──
def main():
    # Read input
    if len(sys.argv) > 1 and os.path.isfile(sys.argv[1]):
        with open(sys.argv[1], 'r', encoding='utf-8') as f:
            raw_text = f.read()
    elif not sys.stdin.isatty():
        raw_text = sys.stdin.read()
    else:
        print("Usage: python cdrno_generate_docx.py [input.fasta]  OR  pipe FASTA via stdin")
        print("       echo \">M1 VH\\nEVQLLES...\" | python cdrno_generate_docx.py")
        sys.exit(1)

    parsed = parse_fasta(raw_text)
    if not parsed:
        print("ERROR: No valid sequences found in input.")
        sys.exit(1)

    # Step 1: Detect chains, validate
    detected = []
    mismatches = []
    for raw_label, label, seq in parsed:
        ct = detect_chain(seq)
        detected.append((label, ct, seq))
        # Check label vs detection
        parts = label.split()
        if parts:
            last = parts[-1].upper()
            if last == 'H' and ct != 'VH':
                mismatches.append((label, 'VH', ct))
            elif last == 'L' and ct not in ('VK', 'VL'):
                mismatches.append((label, 'VL', ct))

    # Step 2: Analyze CDRs & FRs
    cdr_raw = []
    fr_raw = []
    for label, ct, seq in detected:
        for d in SCHEMES:
            r = analyze_cdr(seq, d)
            for reg in r['regions']:
                if reg['name'].startswith('CDR'):
                    cdr_raw.append((label, ct, reg['name'], d, reg['sequence']))
                elif reg['name'].startswith('FR'):
                    fr_raw.append((label, ct, reg['name'], d, reg['sequence']))

    # Step 3: Assign SEQ ID NOs
    vseq_to_id = {}
    nid = 1
    for label, ct, seq in detected:
        if seq not in vseq_to_id:
            vseq_to_id[seq] = nid
            nid += 1

    cdr_seq_to_id = {}
    for lb, ct, cn, sc, seq in cdr_raw:
        if seq not in cdr_seq_to_id:
            cdr_seq_to_id[seq] = nid
            nid += 1

    fr_seq_to_id = {}
    for lb, ct, fn, sc, seq in fr_raw:
        if seq in vseq_to_id:
            fr_seq_to_id[seq] = vseq_to_id[seq]
        elif seq in cdr_seq_to_id:
            fr_seq_to_id[seq] = cdr_seq_to_id[seq]
        elif seq not in fr_seq_to_id:
            fr_seq_to_id[seq] = nid
            nid += 1

    # Build all_unique
    vmeta = {}
    for l, ct, s in detected:
        sid = vseq_to_id[s]
        vmeta.setdefault(s, (sid, []))
        vmeta[s][1].append(f'{l} [VH]' if ct == 'VH' else f'{l} [VL]')

    all_unique = []
    for s, (sid, ds) in vmeta.items():
        all_unique.append((sid, 'V-region', s, ', '.join(ds)))

    for seq, sid in sorted(cdr_seq_to_id.items(), key=lambda x: x[1]):
        srcs = []; seen = set(); ctype = 'CDR'
        for lb, ct, cn, sc, seq2 in cdr_raw:
            if seq2 == seq:
                ctype = cn
                k = (lb, sc)
                if k not in seen:
                    srcs.append(f'{lb} ({sc})')
                    seen.add(k)
        all_unique.append((sid, ctype, seq, ', '.join(srcs)))

    for seq, sid in sorted(fr_seq_to_id.items(), key=lambda x: x[1]):
        srcs = []; seen = set(); ftype = 'FR'
        for lb, ct, fn, sc, seq2 in fr_raw:
            if seq2 == seq:
                ftype = fn
                k = (lb, sc)
                if k not in seen:
                    srcs.append(f'{lb} ({sc})')
                    seen.add(k)
        fr_desc = ', '.join(srcs)
        existing = [e for e in all_unique if e[0] == sid]
        if existing:
            existing[0] = (sid, existing[0][1], existing[0][2],
                           existing[0][3] + f'; [{ftype}] {fr_desc}')
        else:
            all_unique.append((sid, ftype, seq, fr_desc))

    all_unique.sort(key=lambda x: x[0])

    # ── Build DOCX ──
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(9)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Title
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = tp.add_run('CDRNO Analysis — SEQ ID NO Sequence List')
    rr.bold = True; rr.font.size = Pt(14)

    vh_n = sum(1 for _, ct, _ in detected if ct == 'VH')
    vl_n = sum(1 for _, ct, _ in detected if ct in ('VK', 'VL'))
    ip = doc.add_paragraph()
    rr = ip.add_run(
        f'Total unique: {len(all_unique)} (V-region: {len(vmeta)}, CDR: {len(cdr_seq_to_id)}, FR: {len(fr_seq_to_id)})  |  '
        f'Input: {len(detected)} sequences ({vh_n} VH + {vl_n} VL)  |  '
        f'Schemes: Kabat / Chothia / IMGT / AbM'
    )
    rr.font.size = Pt(9)
    doc.add_paragraph()

    # Mismatch warning
    if mismatches:
        wp = doc.add_paragraph()
        rr = wp.add_run('[!] Chain type mismatch (user label vs auto-detected):')
        rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = RGBColor(200, 0, 0)
        for ml, exp, got in mismatches:
            lp = doc.add_paragraph()
            rr = lp.add_run(f'    {ml}: label suggests {exp}, but sequence detected as {got}. Using auto-detected chain for CDR naming.')
            rr.font.size = Pt(8); rr.font.color.rgb = RGBColor(180, 0, 0)
        doc.add_paragraph()

    # ── Table 0: V-region ──
    h0 = doc.add_paragraph()
    rr = h0.add_run('Table 0: V-region Sequences'); rr.bold = True; rr.font.size = Pt(11)

    t0 = doc.add_table(rows=1, cols=5); t0.style = 'Table Grid'
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(t0, ['#', 'User Label', 'Chain', 'Length', 'SEQ ID NO'],
                   [Cm(1.0), Cm(3.5), Cm(2.0), Cm(1.5), Cm(2.5)])

    for i, (l, ct, s) in enumerate(detected):
        row = t0.add_row()
        chain_show = 'VH' if ct == 'VH' else 'VL'
        vals = [str(i+1), l, chain_show, str(len(s)), f'SEQ ID NO: {vseq_to_id[s]}']
        for ci, v in enumerate(vals):
            p = row.cells[ci].paragraphs[0]
            if ci in (0, 3): p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = p.add_run(v); rr.font.size = Pt(8)
            parts = l.split()
            if parts:
                last = parts[-1].upper()
                if (last == 'H' and ct != 'VH') or (last == 'L' and ct not in ('VK', 'VL')):
                    if ci == 2: rr.font.color.rgb = RGBColor(200, 0, 0)
    doc.add_paragraph()

    # ── Table 1: Merged CDR ──
    def build_cdr_rows(chain_type, cdr_names):
        vlist = [(l, ct, s) for l, ct, s in detected if ct == chain_type]
        if not vlist: return []
        ref = vlist[0][0]; rows = []
        for l, ct, s in vlist:
            for cn in cdr_names:
                cells = []
                for d in SCHEMES:
                    fm = [(seq, cdr_seq_to_id.get(seq)) for lb, ct2, cn2, sc, seq in cdr_raw
                          if lb == l and cn2 == cn and sc == d]
                    sid = fm[0][1] if fm else None
                    rf = [(seq2, cdr_seq_to_id.get(seq2)) for lb2, ct3, cn3, sc3, seq2 in cdr_raw
                          if lb2 == ref and cn3 == cn and sc3 == d]
                    rsid = rf[0][1] if rf else None
                    cells.append((sid, l != ref and sid != rsid))
                rows.append((l, cn, cells))
        return rows

    vh_rows = build_cdr_rows('VH', ['CDR-H1', 'CDR-H2', 'CDR-H3'])
    vl_rows = build_cdr_rows('VK', ['CDR-L1', 'CDR-L2', 'CDR-L3']) + \
              build_cdr_rows('VL', ['CDR-L1', 'CDR-L2', 'CDR-L3'])
    all_cdr_rows = vh_rows + vl_rows

    if all_cdr_rows:
        h1 = doc.add_paragraph()
        rr = h1.add_run('Table 1: CDR Annotation — All Chains'); rr.bold = True; rr.font.size = Pt(11)

        t1 = doc.add_table(rows=1, cols=7); t1.style = 'Table Grid'
        t1.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_header_row(t1, ['Variant', 'CDR'] + SCHEMES,
                       [Cm(2.5), Cm(2.0)] + [Cm(3.5)] * 4)

        for l, cn, cells in all_cdr_rows:
            row = t1.add_row()
            p0 = row.cells[0].paragraphs[0]; r0 = p0.add_run(l); r0.font.size = Pt(8)
            p1 = row.cells[1].paragraphs[0]; r1 = p1.add_run(cn); r1.font.size = Pt(8)
            for di, (sid, diff) in enumerate(cells):
                p = row.cells[di+2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f'SEQ ID NO: {sid}' if sid else '—'); r.font.size = Pt(8)
                if diff: r.bold = True; r.font.color.rgb = RGBColor(180, 0, 0)
        doc.add_paragraph()

    # ── Table 3: Merged FR ──
    def build_fr_rows(chain_type, fr_names):
        vlist = [(l, ct, s) for l, ct, s in detected if ct == chain_type]
        if not vlist: return []
        ref = vlist[0][0]; rows = []
        for l, ct, s in vlist:
            for fn in fr_names:
                cells = []
                for d in SCHEMES:
                    fm = [(seq, fr_seq_to_id.get(seq)) for lb, ct2, fn2, sc, seq in fr_raw
                          if lb == l and fn2 == fn and sc == d]
                    sid = fm[0][1] if fm else None
                    rf = [(seq2, fr_seq_to_id.get(seq2)) for lb2, ct3, fn3, sc3, seq2 in fr_raw
                          if lb2 == ref and fn3 == fn and sc3 == d]
                    rsid = rf[0][1] if rf else None
                    cells.append((sid, l != ref and sid != rsid))
                rows.append((l, fn, cells))
        return rows

    vh_fr_rows = build_fr_rows('VH', ['FR1', 'FR2', 'FR3', 'FR4'])
    vl_fr_rows = build_fr_rows('VK', ['FR1', 'FR2', 'FR3', 'FR4']) + \
                 build_fr_rows('VL', ['FR1', 'FR2', 'FR3', 'FR4'])
    all_fr_rows = vh_fr_rows + vl_fr_rows

    if all_fr_rows:
        h3 = doc.add_paragraph()
        rr = h3.add_run('Table 3: FR Annotation — All Chains'); rr.bold = True; rr.font.size = Pt(11)

        t3 = doc.add_table(rows=1, cols=7); t3.style = 'Table Grid'
        t3.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_header_row(t3, ['Variant', 'FR'] + SCHEMES,
                       [Cm(2.5), Cm(2.0)] + [Cm(3.5)] * 4)

        for l, fn, cells in all_fr_rows:
            row = t3.add_row()
            p0 = row.cells[0].paragraphs[0]; r0 = p0.add_run(l); r0.font.size = Pt(8)
            p1 = row.cells[1].paragraphs[0]; r1 = p1.add_run(fn); r1.font.size = Pt(8)
            for di, (sid, diff) in enumerate(cells):
                p = row.cells[di+2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f'SEQ ID NO: {sid}' if sid else '—'); r.font.size = Pt(8)
                if diff: r.bold = True; r.font.color.rgb = RGBColor(180, 0, 0)
        doc.add_paragraph()

    # ── Table 2: Complete SEQ ID NO List ──
    h2 = doc.add_paragraph()
    rr = h2.add_run(f'Table 2: Complete SEQ ID NO Sequence List ({len(all_unique)} entries)')
    rr.bold = True; rr.font.size = Pt(11)

    t2 = doc.add_table(rows=1, cols=4); t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.autofit = False
    add_header_row(t2, ['SEQ ID NO', 'Source Variant(s)', 'Sequence', 'Length'],
                   [Cm(2.2), Cm(7.0), Cm(15.5), Cm(1.3)])

    for sid, typ, seq, desc in all_unique:
        row = t2.add_row()
        for ci, w in enumerate([Cm(2.2), Cm(7.0), Cm(15.5), Cm(1.3)]):
            row.cells[ci].width = w

        source_text = f'[{typ}] {desc}'
        vals = [str(sid), source_text, seq, str(len(seq))]
        for ci, v in enumerate(vals):
            p = row.cells[ci].paragraphs[0]
            if ci in (0, 3): p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ci == 2:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                rr = p.add_run(v)
                rr.font.name = 'Consolas'; rr.font.size = Pt(7)
            elif ci == 1:
                rr = p.add_run(v); rr.font.size = Pt(7.5)
            else:
                rr = p.add_run(v); rr.font.size = Pt(8)

        if typ == 'V-region':
            for ci in range(4):
                s_elm = row.cells[ci]._element.get_or_add_tcPr()
                s_elm.append(s_elm.makeelement(qn('w:shd'), {qn('w:fill'): 'D6E4F0', qn('w:val'): 'clear'}))
        elif typ.startswith('FR'):
            for ci in range(4):
                s_elm = row.cells[ci]._element.get_or_add_tcPr()
                s_elm.append(s_elm.makeelement(qn('w:shd'), {qn('w:fill'): 'E2EFDA', qn('w:val'): 'clear'}))

    # ── Patent Specification Text ──
    # Build lookups
    seq2id = {}  # sequence → SEQ ID NO
    for sid, typ, seq, desc in all_unique:
        if seq not in seq2id:
            seq2id[seq] = sid

    def get_cdr_sid(label, cdr_name, scheme):
        """获取指定变体+CDR+方案的 SEQ ID NO"""
        for lb, ct, cn, sc, seq in cdr_raw:
            if lb == label and cn == cdr_name and sc == scheme:
                return seq2id.get(seq)
        return None

    # Separate VH and VL labels
    vh_labels = [l for l, ct, s in detected if ct == 'VH']
    vl_labels = [l for l, ct, s in detected if ct in ('VK', 'VL')]

    # ── Function: Build CDR profile paragraphs ──
    def build_cdr_profile_text(doc, scheme, target_name='X'):
        """生成 Kabat/IMGT CDR 定义变体的专利文字"""
        cdr_names_h = ['CDR-H1', 'CDR-H2', 'CDR-H3']
        cdr_names_l = ['CDR-L1', 'CDR-L2', 'CDR-L3']

        # Collect CDR profiles per variant
        profiles = {}  # key: (h1,h2,h3,l1,l2,l3) tuple of SEQ ID NOs → list of variant labels
        for l, ct, s in detected:
            profile = tuple(get_cdr_sid(l, cn, scheme) for cn in cdr_names_h + cdr_names_l)
            profiles.setdefault(profile, []).append(l)

        if not profiles:
            return

        h = doc.add_paragraph()
        r = h.add_run(f'CDR 定义实施方案（{scheme} 规则）')
        r.bold = True; r.font.size = Pt(11)

        intro = doc.add_paragraph()
        r = intro.add_run(
            f'本发明提供了特异性结合{target_name}的抗{target_name}抗体及其抗原结合片段，其包含：')
        r.font.size = Pt(10)

        unique_profiles = list(profiles.items())
        for idx, (profile, labels) in enumerate(unique_profiles):
            h1, h2, h3, l1, l2, l3 = profile
            p = doc.add_paragraph()
            r = p.add_run(
                f'{idx+1}）包含如SEQ ID NO: {h1}、{h2}和{h3}所示序列，或相对于所述序列含有一个或多个且'
                f'不超过3个氨基酸的氨基酸取代(例如保守性取代)、缺失或插入的序列的HCDR1、HCDR2、HCDR3；'
                f'和如SEQ ID NO: {l1}、{l2}和{l3}所示序列，或相对于所述序列含有一个或多个且'
                f'不超过3个氨基酸的氨基酸取代(例如保守性取代)、缺失或插入的序列的LCDR1、LCDR2、LCDR3；')
            r.font.size = Pt(10)

        closer = doc.add_paragraph()
        r = closer.add_run(
            f'其中，1）~{len(unique_profiles)}）所述的CDR氨基酸序列是按照{scheme}规则定义的。')
        r.font.size = Pt(10)
        doc.add_paragraph()

    # ── Function: Build V-region overview ──
    def build_vregion_overview(doc, target_name='X'):
        h = doc.add_paragraph()
        r = h.add_run('重链可变区和轻链可变区（全景覆盖）')
        r.bold = True; r.font.size = Pt(11)

        p = doc.add_paragraph()
        vh_ids = sorted(set(vseq_to_id[s] for l, ct, s in detected if ct == 'VH'))
        vl_ids = sorted(set(vseq_to_id[s] for l, ct, s in detected if ct in ('VK', 'VL')))
        vh_str = '、'.join(str(x) for x in vh_ids)
        vl_str = '、'.join(str(x) for x in vl_ids)

        r = p.add_run(
            f'本发明提供了特异性结合{target_name}的抗{target_name}抗体及其抗原结合片段，'
            f'其包含重链可变区和轻链可变区，其中：'
            f'所述重链可变区包含SEQ ID NO：{vh_str}任一项所示的重链可变区包含的HCDR1、HCDR2和HCDR3，'
            f'所述轻链可变区包含SEQ ID NO：{vl_str}任一项所示的轻链可变区包含的LCDR1、LCDR2和LCDR3。')
        r.font.size = Pt(10)

        note = doc.add_paragraph()
        r = note.add_run('在一个实施方案中，上述CDR按照Kabat、IMGT、AbM或Chothia规则定义。')
        r.font.size = Pt(10)
        doc.add_paragraph()

    # ── Function: Build individual VH listing ──
    def build_vh_listing(doc, target_name='X'):
        h = doc.add_paragraph()
        r = h.add_run('重链可变区')
        r.bold = True; r.font.size = Pt(11)

        intro = doc.add_paragraph()
        r = intro.add_run(
            f'在一个实施方案中，本发明提供了特异性结合{target_name}的抗{target_name}抗体及其抗原结合片段，'
            f'其包含重链可变区，其中：')
        r.font.size = Pt(10)

        unique_vh = {}
        for l, ct, s in detected:
            if ct == 'VH':
                unique_vh.setdefault(s, []).append(l)
        for idx, (seq, labels) in enumerate(unique_vh.items()):
            sid = vseq_to_id[seq]
            p = doc.add_paragraph()
            r = p.add_run(
                f'{idx+1}）所述重链可变区包含如SEQ ID NO：{sid}所示氨基酸序列，'
                f'或与SEQ ID NO：{sid}的氨基酸序列具有至少90%、91%、92%、93%、94%、95%、96%、97%、98%'
                f'或99%同一性的氨基酸序列，或由SEQ ID NO：{sid}组成；')
            r.font.size = Pt(10)
        doc.add_paragraph()

    # ── Function: Build individual VL listing ──
    def build_vl_listing(doc, target_name='X'):
        h = doc.add_paragraph()
        r = h.add_run('轻链可变区')
        r.bold = True; r.font.size = Pt(11)

        intro = doc.add_paragraph()
        r = intro.add_run(
            f'在一个实施方案中，本发明提供了特异性结合{target_name}的抗{target_name}抗体及其抗原结合片段，'
            f'其包含轻链可变区，其中：')
        r.font.size = Pt(10)

        unique_vl = {}
        for l, ct, s in detected:
            if ct in ('VK', 'VL'):
                unique_vl.setdefault(s, []).append(l)
        for idx, (seq, labels) in enumerate(unique_vl.items()):
            sid = vseq_to_id[seq]
            p = doc.add_paragraph()
            r = p.add_run(
                f'{idx+1}）所述轻链可变区包含如SEQ ID NO：{sid}所示氨基酸序列，'
                f'或与SEQ ID NO：{sid}的氨基酸序列具有至少90%、91%、92%、93%、94%、95%、96%、97%、98%'
                f'或99%同一性的氨基酸序列，或由SEQ ID NO：{sid}组成；')
            r.font.size = Pt(10)
        doc.add_paragraph()

    # ── Function: Build VH+VL pairing listing ──
    def build_pairing_listing(doc, target_name='X'):
        h = doc.add_paragraph()
        r = h.add_run('重链可变区与轻链可变区配对组合')
        r.bold = True; r.font.size = Pt(11)

        intro = doc.add_paragraph()
        r = intro.add_run(
            f'在另一个实施方案中，本发明提供了特异性结合{target_name}的抗{target_name}抗体及其抗原结合片段，'
            f'其包含重链可变区和轻链可变区，其中：')
        r.font.size = Pt(10)

        # Match VH and VL from same variant label prefix
        # Extract base name (e.g. "M1 VH" → "M1")
        vh_map = {}  # base_name → (label, seq)
        vl_map = {}  # base_name → (label, seq)
        for l, ct, s in detected:
            parts = l.split()
            base = parts[0] if parts else l
            if ct == 'VH':
                vh_map[base] = (l, s)
            elif ct in ('VK', 'VL'):
                vl_map[base] = (l, s)

        idx = 0
        for base in sorted(set(list(vh_map.keys()) + list(vl_map.keys()))):
            vh = vh_map.get(base)
            vl = vl_map.get(base)
            if vh and vl:
                idx += 1
                vh_sid = vseq_to_id[vh[1]]
                vl_sid = vseq_to_id[vl[1]]
                p = doc.add_paragraph()
                r = p.add_run(
                    f'{idx}）所述重链可变区包含如SEQ ID NO：{vh_sid}所示氨基酸序列，'
                    f'或与SEQ ID NO：{vh_sid}的氨基酸序列具有至少90%、91%、92%、93%、94%、95%、96%、97%、98%'
                    f'或99%同一性的氨基酸序列，或由SEQ ID NO：{vh_sid}组成，'
                    f'所述轻链可变区包含如SEQ ID NO：{vl_sid}所示氨基酸序列，'
                    f'或与SEQ ID NO：{vl_sid}的氨基酸序列具有至少90%、91%、92%、93%、94%、95%、96%、97%、98%'
                    f'或99%同一性的氨基酸序列，或由SEQ ID NO：{vl_sid}组成；')
                r.font.size = Pt(10)

        if idx == 0:
            p = doc.add_paragraph()
            r = p.add_run('（根据输入数据中 VH/VL 的对应关系填入配对）')
            r.font.size = Pt(10); r.italic = True
        doc.add_paragraph()

    # ── Generate patent text sections ──
    target = 'X'  # placeholder, edit after generation

    sep = doc.add_paragraph()
    r = sep.add_run('═' * 50)
    r.font.size = Pt(8); r.font.color.rgb = RGBColor(150, 150, 150)

    title_section = doc.add_paragraph()
    r = title_section.add_run('专利申请说明书文字（根据 cdrno 序列表自动生成）')
    r.bold = True; r.font.size = Pt(13)

    build_cdr_profile_text(doc, 'Kabat', target)
    build_cdr_profile_text(doc, 'IMGT', target)
    build_cdr_profile_text(doc, 'AbM', target)
    build_vregion_overview(doc, target)
    build_vh_listing(doc, target)
    build_vl_listing(doc, target)
    build_pairing_listing(doc, target)

    # ── Save ──
    output_path = os.path.join(os.getcwd(), 'CDRNO_SEQ_ID_NO_List.docx')
    doc.save(output_path)
    print(f'Done: {output_path}')
    print(f'V-region: {len(vmeta)} unique | CDR: {len(cdr_seq_to_id)} unique | FR: {len(fr_seq_to_id)} unique | Total: {len(all_unique)}')
    if mismatches:
        print(f'\n[!] Chain mismatches detected:')
        for ml, exp, got in mismatches:
            print(f'    {ml}: label={exp}, detected={got}')

if __name__ == '__main__':
    main()
