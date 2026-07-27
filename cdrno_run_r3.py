#!/usr/bin/env python3
"""CDRNO — R3-15 / R3-3 analysis with VL labels"""
import sys, re
sys.path.insert(0, '.')
from cdr_batch import analyze_cdr, detect_chain
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

seqs_raw = [
    ('R3-15 H', 'QVQLQQSGAELVRPGSSVKISCKASGYAFSSYWMNWVKQRPGQGLEWIGQIYPGDGDARYNGKFNGKATLTTDKSSSTAYMQLSSLTSEDSAVYFCTRSMGLGLDYWGQGTTLTVSS'),
    ('R3-15 L', 'DVVMTQTPLSLPVSLGHQASISCRSSQSLVHSNGNTYLHWYLQKPGQSPKLLIYKVSNRFSGVPDRFSGSGSGTDFTLKISRVEAEDLGVYFCSQTTHVPPYTFGGGTKLEIK'),
    ('R3-3 H',  'QVTLKESGPGILQSSQTLSLTCSFSGFSLSTSGMGVSWIRQPSGKGLEWLAHIYWDDDKRYNPSLKSRLTISKDTSRNQVFLKITSVDTADTATYYCARRAPYFDSPFAYWGQGTLVTVSA'),
    ('R3-3 L',  'NIMMTQSPSSLAVSAGEKVTMSCKSSQSVLYSSNQKNYLAWYQQKPGQSPKLLIYWASTRESGVPDRFTGSGSGTDFTLTISSVQAEDLAVYYCHQYLSSLTFGAGTMLELK'),
]

schemes = ['Kabat', 'Chothia', 'IMGT', 'AbM']

# ── Parse ──
parsed = []
mismatches = []
for label, s in seqs_raw:
    sc = re.sub(r'[^ACDEFGHIKLMNPQRSTVWY]', '', s.upper())
    ct = detect_chain(sc)
    parsed.append((label, ct, sc))
    parts = label.split()
    if parts:
        last = parts[-1].upper()
        if last == 'H' and ct != 'VH':
            mismatches.append((label, 'VH', ct))
        elif last == 'L' and ct not in ('VK', 'VL'):
            mismatches.append((label, 'VL', ct))

# ── Analyze CDRs ──
cdr_raw = []
for label, ct, s in parsed:
    for d in schemes:
        r = analyze_cdr(s, d)
        for reg in r['regions']:
            if reg['name'].startswith('CDR'):
                cdr_raw.append((label, ct, reg['name'], d, reg['sequence']))

# ── SEQ ID NO ──
vseq_to_id = {}; nid = 1
for label, ct, s in parsed:
    if s not in vseq_to_id:
        vseq_to_id[s] = nid
        nid += 1

cdr_seq_to_id = {}
for lb, ct, cn, d, seq in cdr_raw:
    if seq not in cdr_seq_to_id:
        cdr_seq_to_id[seq] = nid
        nid += 1

# ── All unique ──
vmeta = {}
for l, ct, s in parsed:
    sid = vseq_to_id[s]
    vmeta.setdefault(s, (sid, []))
    vmeta[s][1].append(f'{l} [VL]' if ct in ('VK', 'VL') else f'{l} [VH]')

all_u = []
for s, (sid, ds) in vmeta.items():
    all_u.append((sid, 'V-region', s, ', '.join(ds)))

for seq, sid in sorted(cdr_seq_to_id.items(), key=lambda x: x[1]):
    srcs = []; seen = set(); ctype = 'CDR'
    for lb, ct, cn, sc, seq2 in cdr_raw:
        if seq2 == seq:
            ctype = cn
            k = (lb, sc)
            if k not in seen:
                srcs.append(f'{lb} ({sc})')
                seen.add(k)
    all_u.append((sid, ctype, seq, ', '.join(srcs)))
all_u.sort(key=lambda x: x[0])

# ── DOCX ──
doc = Document()
section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(1.2)
section.right_margin = Cm(1.2)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# Title
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = tp.add_run('CDRNO Analysis \u2014 SEQ ID NO Sequence List')
rr.bold = True
rr.font.size = Pt(14)

vh_count = sum(1 for _, ct, _ in parsed if ct == 'VH')
vl_count = sum(1 for _, ct, _ in parsed if ct in ('VK', 'VL'))
info_text = (
    f'Total unique: {len(all_u)} (V-region: {len(vmeta)}, CDR: {len(cdr_seq_to_id)})  |  '
    f'Input: {len(parsed)} sequences ({vh_count} VH + {vl_count} VL)  |  '
    f'Schemes: Kabat / Chothia / IMGT / AbM'
)
ip = doc.add_paragraph()
rr = ip.add_run(info_text)
rr.font.size = Pt(9)
doc.add_paragraph()

# Mismatch warning
if mismatches:
    wp = doc.add_paragraph()
    rr = wp.add_run('[!] Chain type mismatch (user label vs auto-detected):')
    rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = RGBColor(200, 0, 0)
    for ml, exp, got in mismatches:
        lp = doc.add_paragraph()
        rr = lp.add_run(f'    {ml}: label suggests {exp}, detected as {got}. Using auto-detected chain for CDR naming.')
        rr.font.size = Pt(8); rr.font.color.rgb = RGBColor(180, 0, 0)
    doc.add_paragraph()

# Helper
def add_hdr(tbl, texts, widths):
    for i, (txt, w) in enumerate(zip(texts, widths)):
        tbl.rows[0].cells[i].width = w
        p = tbl.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(255, 255, 255)
        shd = tbl.rows[0].cells[i]._element.get_or_add_tcPr()
        shd.append(shd.makeelement(qn('w:shd'), {qn('w:fill'): '2F5496', qn('w:val'): 'clear'}))

# ── Table 0: V-region ──
t0h = doc.add_paragraph()
rr = t0h.add_run('Table 0: V-region Sequences')
rr.bold = True; rr.font.size = Pt(11)

t0 = doc.add_table(rows=1, cols=5)
t0.style = 'Table Grid'
t0.alignment = WD_TABLE_ALIGNMENT.CENTER
add_hdr(t0, ['#', 'User Label', 'Chain', 'Length', 'SEQ ID NO'],
        [Cm(1.0), Cm(3.5), Cm(2.0), Cm(1.5), Cm(2.5)])

for i, (l, ct, s) in enumerate(parsed):
    row = t0.add_row()
    chain_show = 'VH' if ct == 'VH' else 'VL'
    vals = [str(i+1), l, chain_show, str(len(s)), f'SEQ ID NO: {vseq_to_id[s]}']
    for ci, v in enumerate(vals):
        p = row.cells[ci].paragraphs[0]
        if ci in (0, 3):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(v)
        rr.font.size = Pt(8)
        # Red for mismatch
        parts = l.split()
        if parts:
            last = parts[-1].upper()
            if (last == 'H' and ct != 'VH') or (last == 'L' and ct not in ('VK', 'VL')):
                if ci == 2:
                    rr.font.color.rgb = RGBColor(200, 0, 0)

doc.add_paragraph()

# ── Table 1: CDR ──
def build_rows(chain_type, cdr_names):
    vlist = [(l, ct, s) for l, ct, s in parsed if ct == chain_type]
    if not vlist:
        return []
    ref = vlist[0][0]
    rows = []
    for l, ct, s in vlist:
        for cn in cdr_names:
            cells = []
            for d in schemes:
                fm = [(seq, cdr_seq_to_id.get(seq))
                      for lb, ct2, cn2, sc, seq in cdr_raw
                      if lb == l and cn2 == cn and sc == d]
                sid = fm[0][1] if fm else None
                rf = [(seq2, cdr_seq_to_id.get(seq2))
                      for lb2, ct3, cn3, sc3, seq2 in cdr_raw
                      if lb2 == ref and cn3 == cn and sc3 == d]
                rsid = rf[0][1] if rf else None
                cells.append((sid, l != ref and sid != rsid))
            rows.append((l, cn, cells))
    return rows

vh_rows = build_rows('VH', ['CDR-H1', 'CDR-H2', 'CDR-H3'])
vl_rows = build_rows('VK', ['CDR-L1', 'CDR-L2', 'CDR-L3']) + \
          build_rows('VL', ['CDR-L1', 'CDR-L2', 'CDR-L3'])

# Merge VH + VL into one Table 1
all_cdr_rows = vh_rows + vl_rows
if all_cdr_rows:
    hd = doc.add_paragraph()
    rr = hd.add_run('Table 1: CDR Annotation \u2014 All Chains')
    rr.bold = True
    rr.font.size = Pt(11)

    t1 = doc.add_table(rows=1, cols=7)
    t1.style = 'Table Grid'
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_hdr(t1, ['Variant', 'CDR'] + schemes,
            [Cm(2.5), Cm(2.0)] + [Cm(3.5)] * 5)

    for l, cn, cells in all_cdr_rows:
        row = t1.add_row()
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(l); r0.font.size = Pt(8)
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(cn); r1.font.size = Pt(8)
        for di, (sid, diff) in enumerate(cells):
            p = row.cells[di+2].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'SEQ ID NO: {sid}' if sid else '\u2014')
            r.font.size = Pt(8)
            if diff:
                r.bold = True
                r.font.color.rgb = RGBColor(180, 0, 0)
    doc.add_paragraph()

# ── Table 2: Complete SEQ ID NO List ──
hd2 = doc.add_paragraph()
rr = hd2.add_run(f'Table 2: Complete SEQ ID NO Sequence List ({len(all_u)} entries)')
rr.bold = True
rr.font.size = Pt(11)

t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
# Columns: SEQ ID NO | Source Variant(s) | Sequence | Length
t2.autofit = False
add_hdr(t2, ['SEQ ID NO', 'Source Variant(s)', 'Sequence', 'Length'],
        [Cm(2.2), Cm(7.0), Cm(15.5), Cm(1.3)])

for sid, typ, seq, desc in all_u:
    row = t2.add_row()
    # Set row height to auto (no fixed height)
    # Set fixed widths on each cell
    for ci, w in enumerate([Cm(2.2), Cm(7.0), Cm(15.5), Cm(1.3)]):
        row.cells[ci].width = w

    # Source Variant(s) — includes type prefix
    source_text = f'[{typ}] {desc}'

    vals = [str(sid), source_text, seq, str(len(seq))]
    for ci, v in enumerate(vals):
        p = row.cells[ci].paragraphs[0]
        if ci in (0, 3):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ci == 2:
            # Sequence — monospace, small font, allow wrap
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            rr = p.add_run(v)
            rr.font.name = 'Consolas'
            rr.font.size = Pt(7)
        elif ci == 1:
            rr = p.add_run(v)
            rr.font.size = Pt(7.5)
        else:
            rr = p.add_run(v)
            rr.font.size = Pt(8)

    # Blue background for V-region
    if typ == 'V-region':
        for ci in range(4):
            s_elm = row.cells[ci]._element.get_or_add_tcPr()
            s_elm.append(s_elm.makeelement(qn('w:shd'), {qn('w:fill'): 'D6E4F0', qn('w:val'): 'clear'}))

# ── Save ──
output_path = r'C:\Users\admin\Desktop\CC\work\CDRNO_SEQ_ID_NO_List.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
print(f'V-region: {len(vmeta)} | CDR: {len(cdr_seq_to_id)} | Total: {len(all_u)}')
if mismatches:
    print(f'Mismatches: {mismatches}')
